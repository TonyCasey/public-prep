"""File processing service.

Handles file validation and text extraction from uploaded documents.
"""

import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# Supported file extensions
SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx", ".txt"}

# Maximum file size in bytes (5MB)
MAX_FILE_SIZE = 5 * 1024 * 1024


def validate_file_type(filename: str) -> bool:
    """Validate file extension.

    Args:
        filename: Original filename

    Returns:
        True if file type is supported
    """
    ext = Path(filename).suffix.lower()
    return ext in SUPPORTED_EXTENSIONS


def validate_file_size(content: bytes) -> bool:
    """Validate file size.

    Args:
        content: File content as bytes

    Returns:
        True if file size is within limit
    """
    return len(content) <= MAX_FILE_SIZE


async def extract_text_from_file(content: bytes, filename: str) -> str:
    """Extract text content from uploaded file.

    Supports PDF, DOC, DOCX, and TXT files.

    Args:
        content: File content as bytes
        filename: Original filename (used to determine file type)

    Returns:
        Extracted text content

    Raises:
        ValueError: If file type is not supported or extraction fails
    """
    ext = Path(filename).suffix.lower()

    if ext == ".txt":
        return _extract_from_txt(content)
    elif ext == ".pdf":
        return await _extract_from_pdf(content)
    elif ext in {".doc", ".docx"}:
        return await _extract_from_docx(content)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_from_txt(content: bytes) -> str:
    """Extract text from TXT file.

    Args:
        content: File content as bytes

    Returns:
        Text content
    """
    # Try different encodings
    encodings = ["utf-8", "latin-1", "cp1252"]

    for encoding in encodings:
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue

    # Fallback: decode with errors ignored
    return content.decode("utf-8", errors="ignore")


async def _extract_from_pdf(content: bytes) -> str:
    """Extract text from PDF file.

    Args:
        content: File content as bytes

    Returns:
        Extracted text

    Raises:
        ValueError: If PDF extraction fails
    """
    try:
        # Try using pypdf (lightweight)
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        text_parts = []

        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

        if not text_parts:
            raise ValueError("No text content found in PDF")

        return "\n".join(text_parts)

    except ImportError:
        logger.warning("pypdf not installed, trying pdfplumber")

        try:
            import pdfplumber

            with pdfplumber.open(io.BytesIO(content)) as pdf:
                text_parts = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)

            if not text_parts:
                raise ValueError("No text content found in PDF")

            return "\n".join(text_parts)

        except ImportError:
            raise ValueError(
                "PDF extraction requires pypdf or pdfplumber. "
                "Install with: pip install pypdf"
            )

    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise ValueError(f"Failed to extract text from PDF: {e}")


async def _extract_from_docx(content: bytes) -> str:
    """Extract text from DOC/DOCX file.

    Args:
        content: File content as bytes

    Returns:
        Extracted text

    Raises:
        ValueError: If DOCX extraction fails
    """
    try:
        from docx import Document

        doc = Document(io.BytesIO(content))
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        if not text_parts:
            raise ValueError("No text content found in document")

        return "\n".join(text_parts)

    except ImportError:
        raise ValueError(
            "DOCX extraction requires python-docx. "
            "Install with: pip install python-docx"
        )

    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise ValueError(f"Failed to extract text from document: {e}")
